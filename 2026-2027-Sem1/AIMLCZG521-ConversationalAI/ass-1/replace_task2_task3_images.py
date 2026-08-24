from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


REPORT = Path("PS2_Group129_Report.docx")
UPDATED_REPORT = Path("PS2_Group129_Report_updated.docx")
IMAGE_DIR = Path("Images")


def replace_image_before_caption(document, caption, image_path, width):
    old_image = caption._p.getprevious()
    if old_image is None or not old_image.xpath(
        './/*[local-name()="drawing"]'
    ):
        raise RuntimeError(f"Image before caption not found: {caption.text}")

    new_image = document.add_paragraph()
    new_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_image.add_run().add_picture(str(image_path), width=Inches(width))

    parent = caption._p.getparent()
    parent.remove(old_image)
    parent.insert(parent.index(caption._p), new_image._p)


def add_image_before(document, anchor, image_path, caption_text, width):
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(width))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption_text)
    caption_run.font.size = Pt(9)

    anchor_paragraph = anchor._p if hasattr(anchor, "_p") else anchor
    anchor_paragraph.addprevious(caption_paragraph._p)
    caption_paragraph._p.addprevious(image_paragraph._p)


document = Document(REPORT)
task2_caption = next(
    paragraph
    for paragraph in document.paragraphs
    if paragraph.text.startswith("Task 2 execution evidence:")
)
task3_setup_caption = next(
    paragraph
    for paragraph in document.paragraphs
    if paragraph.text.startswith("Task 3 execution evidence: metric definitions")
)
bge_caption = next(
    paragraph
    for paragraph in document.paragraphs
    if paragraph.text.startswith("Task 3 execution evidence: BGE")
)
distilbert_caption = next(
    paragraph
    for paragraph in document.paragraphs
    if paragraph.text.startswith("Task 3 execution evidence: DistilBERT")
)

task2_caption.text = (
    "Task 2 Virtual Lab execution evidence: CPU timing and embedding shapes."
)
replace_image_before_caption(
    document, task2_caption, IMAGE_DIR / "task2_virtual_lab.png", 6.4
)

task3_setup_caption.text = (
    "Task 3 Virtual Lab execution evidence: metric definitions and pair selection."
)
replace_image_before_caption(
    document, task3_setup_caption, IMAGE_DIR / "task3_Virtual_lab1.png", 6.4
)

add_image_before(
    document,
    bge_caption._p.getprevious(),
    IMAGE_DIR / "task3_Virtual_lab2.png",
    "Task 3 Virtual Lab execution evidence: 25 query-document pairs and similarity values.",
    6.2,
)

bge_caption.text = (
    "Task 3 Virtual Lab execution evidence: BGE ranking and normalization checks."
)
replace_image_before_caption(
    document, bge_caption, IMAGE_DIR / "task3_viratual_lab_3.png", 6.2
)

distilbert_caption.text = (
    "Task 3 Virtual Lab execution evidence: DistilBERT normalization and ranking checks."
)
replace_image_before_caption(
    document, distilbert_caption, IMAGE_DIR / "task3_virual_Lab4.png", 6.2
)

document.save(UPDATED_REPORT)
UPDATED_REPORT.replace(REPORT)