from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


REPORT = Path("PS2_Group129_Report.docx")
UPDATED_REPORT = Path("PS2_Group129_Report_with_task7_image6.docx")
IMAGE_PATH = Path("Images/task7_virtual_lab6.png")


document = Document(REPORT)
task8_heading = next(
    paragraph
    for paragraph in document.paragraphs
    if paragraph.text.startswith("Task 8 — Final Recommendation")
)

image_paragraph = document.add_paragraph()
image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
image_paragraph.add_run().add_picture(str(IMAGE_PATH), width=Inches(6.4))

caption_paragraph = document.add_paragraph()
caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption_run = caption_paragraph.add_run(
    "Task 7 Virtual Lab execution evidence: aggregate cross-model comparison metrics for 300 queries."
)
caption_run.font.size = Pt(9)

task8_heading._p.addprevious(caption_paragraph._p)
caption_paragraph._p.addprevious(image_paragraph._p)

document.save(UPDATED_REPORT)
UPDATED_REPORT.replace(REPORT)